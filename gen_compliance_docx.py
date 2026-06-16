# -*- coding: utf-8 -*-
"""One-page Word version of the RT compliance assessment."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

ACCENT = RGBColor(0x2F, 0x48, 0x58)
GRAY = RGBColor(0x66, 0x6C, 0x77)

doc = Document()
sec = doc.sections[0]
sec.page_height, sec.page_width = Cm(29.7), Cm(21.0)
for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
    setattr(sec, m, Cm(1.3))

style = doc.styles["Normal"]
style.font.name = "Malgun Gothic"
style.font.size = Pt(8.5)
style.element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
style.paragraph_format.space_after = Pt(2)


def para(text, size=8.5, bold=False, color=None, align=None, after=2):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    if color:
        r.font.color.rgb = color
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    return p


def fill(cell, text, size=7.5, bold=False, color=None):
    cell.paragraphs[0].text = ""
    r = cell.paragraphs[0].add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    if color:
        r.font.color.rgb = color
    cell.paragraphs[0].paragraph_format.space_after = Pt(0)


def make_table(headers, rows, widths):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, (htxt, w) in enumerate(zip(headers, widths)):
        fill(t.rows[0].cells[j], htxt, bold=True)
        for i in range(len(rows) + 1):
            t.rows[i].cells[j].width = Cm(w)
    for i, row in enumerate(rows):
        for j, txt in enumerate(row):
            bold = j == len(row) - 1 and txt.startswith("적합")
            fill(t.rows[i + 1].cells[j], txt, bold=bold)
    return t


# ---- header ----
para("방사선투과시험(RT) 기법 적합성 평가서", 14, True, ACCENT,
     WD_ALIGN_PARAGRAPH.CENTER, 0)
para("Radiographic Technique Compliance Assessment — LR Rules / ISO / ASNT·ASME",
     8, False, GRAY, WD_ALIGN_PARAGRAPH.CENTER, 4)

make_table(
    ["대상 용접부", "시스템", "기법", "작성일"],
    [["분기 아웃렛–WN 플랜지 거스 용접부\n(냉간성형 일체형 outlet, OD 26.7 / t 2.87)",
      "IGF Code 가스연료 이중관\n(외관 NPS 3 Sch 10S / 내관 NPS 2 Sch 40)",
      "감마선 RT, DWDI\nIr-192, 중첩기법", "2026-06-12"]],
    [5.0, 5.6, 4.2, 3.0])

para("1. 기법 요약 (Technique Data)", 9.5, True, ACCENT, after=1)
make_table(
    ["SFD / SOD / b", "고도각", "촬영", "필름", "IQI", "Ug"],
    [["210 / 173 / 37 mm", "≈6°\n(플랜지 틈 8 mm 투과)",
      "3회\n(−60°/0°/+60°)", "곡면 R24 스탠드오프\n내관·플랜지 테이프 고정",
      "ISO 19232-1 와이어형\nW10–16, 선원측", "0.43 mm"]],
    [3.2, 2.9, 2.6, 3.4, 3.6, 2.1])
para("커버리지: 회당 필름측 상 + 선원측 상 동시 기록(DWDI) → 60° 간격 3회 합산 시 "
     "용접 전둘레 360° 판독 범위 확보", 7.5, color=GRAY, after=4)

para("2. 적합성 매트릭스", 9.5, True, ACCENT, after=1)
make_table(
    ["규격", "요구사항", "요구치 / 근거", "본 세팅", "판정"],
    [
        ["LR Rules", "가스연료 배관 용접부 100% 체적검사",
         "IGF Code Ch.16.4 이행", "전둘레 3회 촬영, 360% 커버", "적합"],
        ["LR Rules", "RT 기법 규격 / 용접 품질등급",
         "ISO 17636-1 / ISO 5817 Level B",
         "ISO 17636-1 기법, ISO 10675-1 AL-1 판정", "적합"],
        ["LR Rules", "NDE 요원 자격", "ISO 9712(또는 동등) Level 2",
         "검사·판독 Level 2 지정", "적합(조건)"],
        ["ISO 17636-1 §7.6", "최소 SOD (Class A)",
         "f ≥ 7.5·d·b^⅔ → 167 mm", "SOD 173 mm", "적합"],
        ["ISO 17636-1 §7.7", "DWDI 중첩기법 최소 촬영수",
         "3회 (60° 또는 120° 간격)", "3회 @ 60°", "적합"],
        ["ISO 17636-1 §7.8", "IQI 종류·배치",
         "와이어형, 선원측, 용접선 직각", "와이어형 W10–16, 선원측 래핑", "적합"],
        ["ISO 19232-1 Tab.B.1", "요구 식별 와이어 (w≈5.7 mm)",
         "W14 (Ø0.16 mm) 이상", "모의필름 W13–14 수준, 시행 시 실측", "적합(조건)"],
        ["ISO 17636-1 §7.9", "필름 농도", "≥ 2.0 (Class A)",
         "노출조건으로 달성, 농도계 확인", "적합(조건)"],
        ["ISO 17636-1 §7.1.2", "Ir-192 적용두께",
         "권장 w ≥ 10 mm, NOTE: 접근제한 시 합의 적용",
         "w≈5.7 — IQI 감도 입증 + 선급 합의 (Se-75 채택 시 해소)", "합의 필요"],
        ["ASME V T-274.2", "기하학적 불선명도", "Ug ≤ 0.51 mm (t<50.8)",
         "0.43 mm", "적합"],
        ["ASME V T-276/277", "IQI 선정·배치", "와이어형 허용, 선원측 원칙",
         "선원측 와이어형", "적합"],
        ["ASNT SNT-TC-1A", "요원 자격", "Level II 이상 수행·판독",
         "Written Practice에 따라 지정", "적합(조건)"],
    ],
    [2.8, 4.0, 4.6, 4.6, 1.8])

para("비고 — Class B 지정 시: SOD ≥ 333 mm 필요. 선원측은 개방 공간이므로 SFD ≥ 370 mm "
     "연장(Ug ≈ 0.22 mm) 또는 소초점(f=1.0 mm) Ir-192/Se-75 선원 채택으로 충족 가능.",
     7.5, color=GRAY, after=4)

para("3. 결론", 9.5, True, ACCENT, after=1)
para("분기 아웃렛은 내관에서 냉간성형된 일체형(무용접)으로, 체적검사 대상은 아웃렛–"
     "플랜지 거스 용접 1개소이다. 본 RT 세팅은 LR 규칙(ISO 17636-1 Class A 기준), "
     "ASME Sec. V 및 ASNT 요구사항을 충족한다. 시행 단계 확인 항목: ① IQI W14 식별·농도(≥2.0) 실측, "
     "② 요원 자격 증빙(ISO 9712/SNT-TC-1A Lv.2), ③ 박육 Ir-192 적용에 대한 선급 "
     "검사관 합의(또는 Se-75 채택). 판정기준: ISO 5817 Level B / ISO 10675-1 "
     "Acceptance Level 1.", 8)

para("첨부: ① 기하 계산서(rt_calculation.md) ② 3D 배치(viewer_offline.html) "
     "③ 촬영별 배치도·단면도(rt_shot/side 1–3) ④ 모의 투과사진(rt_film 1–3)",
     7, color=GRAY, after=0)

doc.save("RT_Compliance_Assessment.docx")
print("saved RT_Compliance_Assessment.docx")
