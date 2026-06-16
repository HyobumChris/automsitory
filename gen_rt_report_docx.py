# -*- coding: utf-8 -*-
"""RT feasibility & acceptability report (ISO 17636-1 Class A / ASNT·ASME)
with 3D captures, layout diagrams and simulated radiographs."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

ACCENT = RGBColor(0x2F, 0x48, 0x58)
GRAY = RGBColor(0x66, 0x6C, 0x77)

doc = Document()
sec = doc.sections[0]
sec.page_height, sec.page_width = Cm(29.7), Cm(21.0)
for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
    setattr(sec, m, Cm(1.6))

style = doc.styles["Normal"]
style.font.name = "Malgun Gothic"
style.font.size = Pt(9)
style.element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
style.paragraph_format.space_after = Pt(3)


def para(text, size=9, bold=False, color=None, align=None, after=3):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    if color:
        r.font.color.rgb = color
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    return p


def heading(num, text):
    para(f"{num}. {text}", 11, True, ACCENT, after=2)


def caption(text):
    para(text, 7.5, False, GRAY, WD_ALIGN_PARAGRAPH.CENTER, 6)


def fill(cell, text, size=8, bold=False):
    cell.paragraphs[0].text = ""
    r = cell.paragraphs[0].add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    cell.paragraphs[0].paragraph_format.space_after = Pt(0)


def make_table(headers, rows, widths, size=8):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, (htxt, w) in enumerate(zip(headers, widths)):
        fill(t.rows[0].cells[j], htxt, size, True)
        for i in range(len(rows) + 1):
            t.rows[i].cells[j].width = Cm(w)
    for i, row in enumerate(rows):
        for j, txt in enumerate(row):
            bold = j == len(row) - 1 and txt.startswith(("적합", "가능", "합격"))
            fill(t.rows[i + 1].cells[j], txt, size, bold)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return t


def pictures_row(paths, widths):
    """Side-by-side images in a borderless table row."""
    t = doc.add_table(rows=1, cols=len(paths))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, (p, w) in enumerate(zip(paths, widths)):
        cell = t.rows[0].cells[j]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].add_run().add_picture(p, width=Cm(w))


def picture(path, w):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(1)
    p.add_run().add_picture(path, width=Cm(w))


def page_break():
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


# ============================ page 1 ============================
para("분기 거스 용접부 방사선투과시험(RT)", 16, True, ACCENT,
     WD_ALIGN_PARAGRAPH.CENTER, 0)
para("촬영 타당성 및 합격성 평가 보고서", 16, True, ACCENT,
     WD_ALIGN_PARAGRAPH.CENTER, 2)
para("RT Feasibility & Acceptability Report — ISO 17636-1 Class A / "
     "ASNT SNT-TC-1A / ASME Sec. V", 9, False, GRAY,
     WD_ALIGN_PARAGRAPH.CENTER, 8)

make_table(
    ["문서번호", "대상 용접부", "기법", "작성일"],
    [["RT-FR-001", "이중관 분기 아웃렛–WN 플랜지 거스 용접부",
      "감마선 RT · DWDI 중첩기법", "2026-06-12"]],
    [2.6, 8.0, 4.4, 2.6])

heading(1, "목적 및 범위")
para("본 보고서는 IGF Code 적용 가스연료 이중관의 분기 거스 용접부에 대하여, "
     "협소한 환형 공간 조건에서 방사선투과시험(RT)이 물리적으로 수행 가능함을 "
     "기하학적으로 입증하고, 그 기법이 ISO 17636-1 Class A 및 ASNT/ASME 요구사항을 "
     "충족하여 판독 결과가 합격성(acceptability) 평가에 유효함을 기술한다.")

heading(2, "검사 대상")
make_table(
    ["항목", "사양", "항목", "사양"],
    [
        ["외관 (보호관)", "NPS 3 Sch 10S — OD 88.9 / t 3.05",
         "분기 아웃렛", "내관에서 냉간성형(extruded) 일체형 — OD 26.7 / t 2.87"],
        ["내관 (연료관)", "NPS 2 Sch 40 — OD 60.3 / t 3.91",
         "플랜지", "Compact WN — OD 70 / t 14.6, 목 8 mm"],
        ["외관 개구부", "Ø70 원형 컷아웃",
         "대상 용접부", "아웃렛–플랜지 거스 용접 1개소 (무용접 성형 분기루트)"],
        ["용접선 위치", "내관 표면 위 10 mm(플랭크) / 6.85 mm(크라운)",
         "용접선–플랜지 틈", "8 mm"],
    ],
    [2.9, 5.9, 2.9, 5.9], size=7.5)
para("분기 아웃렛은 내관에서 냉간성형으로 일체 성형되어 분기 연결부 자체는 "
     "무용접이며, 체적검사 대상은 아웃렛–플랜지 거스 용접 1개소이다.",
     8, color=GRAY)

pictures_row(["cap_overall.png", "cap_branch.png"], [8.6, 8.6])
caption("그림 1. 검사 대상 — 이중관 전경(좌) 및 분기부 상세(우): 냉간성형 collar, "
        "거스 용접 비드, compact WN 플랜지")

# ============================ page 2 ============================
page_break()
heading(3, "촬영 기법 및 배치")
make_table(
    ["항목", "값", "항목", "값"],
    [
        ["선원", "Ir-192 (초점 d = 2 mm)", "기법", "DWDI 중첩(타원·수직) 기법"],
        ["촬영 횟수", "3회 — 방위각 −60° / 0° / +60°", "선원 고도각",
         "≈6.6° (플랜지 하단 틈 8 mm 투과)"],
        ["SFD / SOD / b", "210 / 173 / 37 mm", "기하학적 불선명도 Ug",
         "0.43 mm"],
        ["필름", "곡면 스트립 R24 (스탠드오프 4 mm), 호 103°, 선원 반대측 — "
         "하단 내관·상단 플랜지 테이프 고정",
         "IQI", "ISO 19232-1 와이어형 W10–16, 선원측 래핑"],
    ],
    [3.2, 5.6, 3.2, 5.6], size=7.5)

picture("cap_rt_all.png", 10.8)
caption("그림 2. 3D 배치 검증 — 3개 선원 위치(빨강)와 가이드 튜브, 분기부를 향한 "
        "빔 경로. 외관 크라운 간섭을 피해 개방 섹터(±z)에 배치")

para("회당 필름측 상과 선원측 상이 동시에 기록(DWDI)되며, 60° 간격 3회 "
     "촬영으로 용접 전둘레 360°의 판독 범위를 확보한다. 각 촬영의 평면 배치, "
     "단면 빔 경로, 예상 투과사진은 그림 3과 같다.", 8.5)

pictures_row(["rt_shot1.png", "rt_side1.png", "rt_film1.png"], [5.5, 5.5, 5.5])
pictures_row(["rt_shot2.png", "rt_side2.png", "rt_film2.png"], [5.5, 5.5, 5.5])
pictures_row(["rt_shot3.png", "rt_side3.png", "rt_film3.png"], [5.5, 5.5, 5.5])
caption("그림 3. Shot 1–3 (위→아래) — 평면 커버리지(좌), 단면 빔 경로(중), "
        "모의 투과사진(우)")

# ============================ page 3 ============================
heading(4, "촬영 가능성(Feasibility) — 기하 검증")
make_table(
    ["검증 항목", "기하 조건", "결과", "판정"],
    [
        ["빔 통과 클리어런스", "선원 고도각 6.6° 중심선의 플랜지 디스크 엣지(R35) 통과고",
         "플랜지 하단보다 약 2 mm 아래로 통과 — 그림자 없음", "가능"],
        ["필름 배치 (스탠드오프)", "용접선이 내관 표면 바로 위(크라운 6.85 mm) — "
         "비드 밀착 시 스트립 좌굴",
         "R24 원통 스탠드오프(4 mm) 채택: 하단은 내관 표면, 상단은 플랜지 디스크 "
         "하면에 테이프 고정", "가능(주1)"],
        ["필름 삽입 공간", "넥 R13.4–컷아웃 R35 환형부, 플랜지 아래 8 mm 틈",
         "R24 스트립 삽입 가능 (상단 여유 2 mm)", "가능"],
        ["IQI 배치", "비드 위 래핑, 슬리브 상단 용접선 +6.5 mm",
         "플랜지 간섭 없음, 선원측 배치 충족", "가능"],
        ["선원 접근", "선원측은 환형부 밖 개방 공간 (반경 185 mm)",
         "가이드 튜브 직선 거치 가능", "가능"],
    ],
    [3.4, 5.6, 6.0, 1.6], size=7.5)
para("주1) 규격 요구는 \"필름을 가능한 한 검사체에 가깝게\"이며 밀착 자체가 "
     "의무는 아니다. 본 건은 용접선–내관 간격이 좁아 밀착 시 스트립이 좌굴되므로 "
     "4 mm 스탠드오프 배치를 채택하고, 그에 따른 b 증가(33→37 mm)를 보상하기 "
     "위해 선원 반경을 170→185 mm로 연장하였다. 결과 Ug 0.43 mm 및 Class A "
     "최소거리(§7.6)를 모두 충족하며(5절), IQI가 선원측 용접부 위에 유지되므로 "
     "스탠드오프를 포함한 전체 기하의 감도가 W14 식별로 입증된다. 3D 모델에 동일 "
     "배치를 반영하였다(그림 4).", 8, color=GRAY)

picture("cap_film_detail.png", 13.5)
caption("그림 4. 필름 스탠드오프 배치 상세 — R24 원통 유지, 하단 내관·상단 플랜지 "
        "테이프 고정, 비드 위 IQI 래핑 (Shot 3만 표시)")

heading(5, "ISO 17636-1 Class A 적합성")
make_table(
    ["조항", "요구사항", "요구치", "본 세팅", "판정"],
    [
        ["§7.6", "최소 선원-시험체 거리", "f ≥ 7.5·d·b^⅔ → 167 mm",
         "SOD 173 mm", "적합"],
        ["§7.7", "DWDI 최소 촬영수", "3회 (60°/120° 간격)", "3회 @ 60°", "적합"],
        ["§7.8", "IQI 종류·배치", "와이어형, 선원측, 용접선 직각",
         "W10–16 선원측 래핑", "적합"],
        ["ISO 19232-1", "요구 식별 와이어 (w≈5.7 mm)", "W14 (Ø0.16) 이상",
         "모의필름 W13–14 식별 수준", "적합(실측)"],
        ["§7.9", "필름 농도", "≥ 2.0", "노출조건으로 달성, 농도계 확인", "적합(실측)"],
        ["§7.1.2", "Ir-192 적용두께", "권장 w ≥ 10 mm (접근제한 시 합의)",
         "w≈5.7 — IQI 감도 입증 + 선급 합의", "합의 필요"],
    ],
    [1.7, 4.2, 4.6, 5.0, 2.1], size=7.5)

# ============================ page 4 ============================
page_break()
heading(6, "ASNT / ASME 적합성")
make_table(
    ["규격", "요구사항", "본 세팅", "판정"],
    [
        ["ASME V T-274.2", "기하학적 불선명도 Ug ≤ 0.51 mm (t < 50.8)",
         "Ug = 0.43 mm", "적합"],
        ["ASME V T-276/277", "IQI 선정·배치 — 와이어형 허용, 선원측 원칙",
         "선원측 와이어형 W10–16", "적합"],
        ["ASME V T-282", "방사선 에너지 적정성", "Ir-192, 투과두께 대비 "
         "IQI 감도로 입증", "적합"],
        ["ASNT SNT-TC-1A", "요원 자격 — Level II 이상 수행·판독",
         "Written Practice에 따라 RT Level II 지정", "적합(증빙)"],
    ],
    [3.0, 6.4, 5.2, 2.0], size=7.5)

heading(7, "모의 투과사진의 합격성(Acceptability) 검토")
para("레이트레이싱 모의 투과사진(그림 3 우측 열)에서 다음을 확인하였다:", 8.5)
make_table(
    ["판독 항목", "모의 결과", "합격 기준 연계"],
    [
        ["DWDI 타원상 분리", "근측·원측 벽 상이 상하로 분리 결상",
         "양 벽 독립 판독 가능 → 100% 체적검사 성립"],
        ["IQI 와이어 식별", "W13–14 수준 식별",
         "ISO 19232-1 요구 W14 충족 (시행 시 실측 확인)"],
        ["플랜지 그림자", "필름 상단 밝은 밴드로 국한",
         "용접부 평가구역과 분리 — 판독 방해 없음"],
        ["판정 기준", "—", "ISO 5817 Level B / ISO 10675-1 Acceptance Level 1"],
    ],
    [3.6, 5.8, 7.2], size=7.5)

heading(8, "결론")
para("① 본 분기 거스 용접부는 용접선–플랜지 8 mm 틈과 환형 개구부를 통한 "
     "필름·선원 배치가 기하학적으로 성립하며, 3회 DWDI 촬영으로 전둘레 100% "
     "체적검사가 가능하다. ② 기법 변수(SOD, Ug, IQI, 촬영수)는 ISO 17636-1 "
     "Class A 및 ASME Sec. V 요구치를 충족하고, ASNT SNT-TC-1A에 따른 Level II "
     "요원 수행을 전제로 판독 결과는 ISO 5817 Level B / ISO 10675-1 AL-1 기준의 "
     "합격성 평가에 유효하다. ③ 시행 단계 확인 항목: IQI W14 식별 및 농도(≥2.0) "
     "실측, 요원 자격 증빙, 박육 Ir-192 적용에 대한 선급 검사관 합의(또는 Se-75 "
     "채택).", 9)

para("첨부: ① 기하 계산서(rt_calculation.md) ② 적합성 평가서(RT_Compliance_"
     "Assessment.docx) ③ 3D 모델(viewer_offline.html) ④ 촬영별 배치도·단면도·"
     "모의필름 원본 PNG", 7.5, color=GRAY, after=10)

t = doc.add_table(rows=2, cols=3)
t.style = "Table Grid"
t.alignment = WD_TABLE_ALIGNMENT.CENTER
for j, h in enumerate(["작성 (RT Level II)", "검토 (RT Level III)", "승인"]):
    fill(t.rows[0].cells[j], h, 8, True)
    t.rows[0].cells[j].width = Cm(5.6)
    fill(t.rows[1].cells[j], "\n\n(서명)          (일자)", 8)
    t.rows[1].cells[j].width = Cm(5.6)

doc.save("RT_Feasibility_Report.docx")
print("saved RT_Feasibility_Report.docx")
